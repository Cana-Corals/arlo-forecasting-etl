"""
Build the thesis supplementary appendix DOCX.

Sections:
  A — Data Dictionary (master dataset + engineered predictors)
  B — Complete Feature Importance Rankings (all four models)
  C — Model Selection and Cross-Validation Detail
  D — Data Pipeline Script Descriptions
  E — External Data Source Documentation
"""

from pathlib import Path
import pandas as pd
import numpy as np

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).resolve().parents[1]
OUTPUTS_DIR = BASE_DIR / "outputs"

# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def set_col_width(cell, width_inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def shade_row(row, hex_color="E8EEF4"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


def style_header_row(row, hex_color="2E4057"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
    return h


def add_table_with_header(doc, headers, data_rows, col_widths=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(font_size)
                run.font.bold = True
    style_header_row(hdr)

    for row_idx, row_data in enumerate(data_rows):
        row = table.add_row()
        if row_idx % 2 == 1:
            shade_row(row, "F5F7FA")
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(font_size)
            if col_widths:
                set_col_width(cell, col_widths[i])

    if col_widths:
        for cell in table.rows[0].cells:
            pass  # widths already set per row above

    return table


# ---------------------------------------------------------------------------
# Data dictionary definitions
# ---------------------------------------------------------------------------

MASTER_VARS = [
    # (variable_name, description, source, units/type)
    # --- Capacity ---
    ("total_physical_rooms",     "Total physical guestrooms in the property", "PMS daily stats (room type)", "Integer"),
    ("ooo_rooms",                "Rooms out of order (maintenance, renovation)", "PMS daily stats", "Integer"),
    ("available_rooms",          "Sellable rooms = total_physical_rooms − ooo_rooms", "Derived", "Integer"),
    # --- Occupancy / revenue ---
    ("room_nights",              "Occupied room nights sold on date", "PMS daily stats", "Integer"),
    ("occupancy_rate",           "room_nights / available_rooms", "Derived", "Decimal 0–1"),
    ("room_revenue",             "Net room revenue collected on stay date", "PMS daily stats", "USD"),
    ("food_revenue",             "Food & beverage revenue attributed to stay date", "PMS daily stats", "USD"),
    ("other_revenue",            "Ancillary revenue (parking, spa, etc.)", "PMS daily stats", "USD"),
    ("total_revenue",            "Sum of room + food + other revenue", "Derived", "USD"),
    ("adr",                      "Average Daily Rate = room_revenue / room_nights", "Derived", "USD / room night"),
    ("revpar",                   "Revenue per Available Room = room_revenue / available_rooms", "Derived", "USD"),
    # --- Operational ---
    ("cancelled_rooms",          "Reservations cancelled with same-day arrival date", "PMS daily stats", "Integer"),
    ("no_show_rooms",            "No-show reservations", "PMS daily stats", "Integer"),
    ("day_use_rooms",            "Day-use bookings (no overnight stay)", "PMS daily stats", "Integer"),
    ("deduct_individual_rooms",  "Individual deduct (complimentary/house) rooms", "PMS daily stats", "Integer"),
    ("deduct_group_rooms",       "Group deduct rooms", "PMS daily stats", "Integer"),
    # --- Pricing ---
    ("retail_rate",              "Best available rack rate published for that date; forward-filled from rate change log", "Rate change XLSX", "USD"),
    # --- Guest satisfaction (Medallia) ---
    ("medallia_sample_size",     "Number of guest survey responses in the reporting week", "Medallia XLS (weekly)", "Integer"),
    ("medallia_overall_satisfaction", "Overall service satisfaction score (10-pt scale)", "Medallia XLS (weekly)", "0–10"),
    ("medallia_likelihood_to_recommend", "Net Promoter proxy — likelihood to recommend score", "Medallia XLS (weekly)", "0–10"),
    ("medallia_likelihood_to_return", "Guest intent to return score", "Medallia XLS (weekly)", "0–10"),
    ("medallia_value_for_price", "Perceived value-for-money score", "Medallia XLS (weekly)", "0–10"),
    ("medallia_hotel_cleanliness","Cleanliness rating", "Medallia XLS (weekly)", "0–10"),
    # --- Calendar ---
    ("day_of_week",              "ISO weekday (0=Monday … 6=Sunday)", "Derived from date", "Integer 0–6"),
    ("month",                    "Calendar month number", "Derived from date", "Integer 1–12"),
    ("quarter",                  "Calendar quarter", "Derived from date", "Integer 1–4"),
    ("week_of_year",             "ISO week number", "Derived from date", "Integer 1–53"),
    ("is_weekend",               "1 if Saturday or Sunday, else 0", "Derived from date", "Boolean"),
    # --- Weather ---
    ("temp_mean_f",              "Daily mean temperature at 2 m, Williamsburg", "Open-Meteo archive API", "°F"),
    ("temp_max_f",               "Daily maximum temperature", "Open-Meteo archive API", "°F"),
    ("temp_min_f",               "Daily minimum temperature", "Open-Meteo archive API", "°F"),
    ("precipitation_in",         "Total daily precipitation (rain + melted snow)", "Open-Meteo archive API", "Inches"),
    ("had_precipitation",        "1 if precipitation_in > 0", "Derived", "Boolean"),
    ("snowfall_in",              "Total daily snowfall", "Open-Meteo archive API", "Inches"),
    ("had_snow",                 "1 if snowfall_in > 0", "Derived", "Boolean"),
    ("windspeed_max_mph",        "Maximum 10-m wind speed during the day", "Open-Meteo archive API", "mph"),
    ("weathercode",              "WMO weather interpretation code (0=clear … 99=violent thunderstorm)", "Open-Meteo archive API", "Integer (WMO)"),
    # --- Holidays ---
    ("is_federal_holiday",       "1 if a US federal holiday, else 0", "Nager.Date API / FederalPay CSV", "Boolean"),
    # --- NYC events ---
    ("event_count",              "Number of flagged events occurring on that date", "League APIs + Ticketmaster", "Integer"),
    ("is_major_event_day",       "1 if any major event (sports, concert, US Open)", "Derived", "Boolean"),
    ("is_barclays_event",        "1 if a Nets home game at Barclays Center", "NBA API (stats.nba.com)", "Boolean"),
    ("is_msg_event",             "1 if a Knicks or Rangers home game at Madison Square Garden", "NBA / NHL API", "Boolean"),
    ("is_yankees_game",          "1 if a Yankees home game at Yankee Stadium", "MLB Stats API", "Boolean"),
    ("is_mets_game",             "1 if a Mets home game at Citi Field", "MLB Stats API", "Boolean"),
    ("is_knicks_game",           "1 if a Knicks home game", "NBA API", "Boolean"),
    ("is_nets_game",             "1 if a Nets home game", "NBA API", "Boolean"),
    ("is_rangers_game",          "1 if a Rangers home game", "NHL API", "Boolean"),
    ("is_nba_game",              "1 if any NBA home game (Knicks or Nets)", "Derived", "Boolean"),
    ("is_mlb_game",              "1 if any MLB home game (Yankees or Mets)", "Derived", "Boolean"),
    ("is_nhl_game",              "1 if any NHL home game (Rangers)", "Derived", "Boolean"),
    ("is_us_open_event",         "1 if within the US Open Tennis tournament window (hardcoded dates, Arthur Ashe)", "Hardcoded", "Boolean"),
    ("is_major_concert",         "1 if a major concert at Barclays, MSG, or Brooklyn Mirage", "Ticketmaster Discovery API", "Boolean"),
    ("is_major_sports_event",    "1 if any major sports event (MLB + NBA + NHL)", "Derived", "Boolean"),
    ("days_to_next_event",       "Days until the next major event, capped at 14", "Derived", "Integer 0–14"),
    # --- STR comp set ---
    ("comp_occ",                 "Comp set occupancy rate (STR weekly report)", "STR data export (XLSX)", "Decimal 0–1"),
    ("comp_adr",                 "Comp set average daily rate", "STR data export (XLSX)", "USD"),
    ("comp_revpar",              "Comp set RevPAR", "STR data export (XLSX)", "USD"),
    ("mpi",                      "Market Penetration Index = Arlo occ / comp occ × 100", "STR data export", "Index (100=parity)"),
    ("ari",                      "Average Rate Index = Arlo ADR / comp ADR × 100", "STR data export", "Index (100=parity)"),
    ("rgi",                      "Revenue Generation Index = Arlo RevPAR / comp RevPAR × 100", "STR data export", "Index (100=parity)"),
]

ENGINEERED_VARS = [
    # Booking pace
    ("pickup_7d",    "Rooms on books booked ≥7 days before arrival (confirmed reservations)", "Derived from res_main_clean.csv", "Integer"),
    ("pickup_14d",   "Rooms on books booked ≥14 days before arrival", "Derived", "Integer"),
    ("pickup_21d",   "Rooms on books booked ≥21 days before arrival", "Derived", "Integer"),
    ("pickup_30d",   "Rooms on books booked ≥30 days before arrival", "Derived", "Integer"),
    ("pickup_60d",   "Rooms on books booked ≥60 days before arrival", "Derived", "Integer"),
    ("total_rooms_on_books", "Total confirmed reservations on the books as of data snapshot date", "Derived from reservations_daily.csv", "Integer"),
    ("avg_booked_rate",      "Mean rate of all confirmed reservations in the books snapshot", "Derived from reservations_daily.csv", "USD"),
    # Holiday proximity
    ("days_to_next_holiday", "Days until the next federal holiday", "Derived from holiday list", "Integer"),
    ("days_from_last_holiday","Days since the previous federal holiday", "Derived from holiday list", "Integer"),
    # Occupancy lags & rolling
    ("occ_lag_7d",   "Occupancy rate 7 days prior", "Lag of occupancy_rate", "Decimal 0–1"),
    ("occ_lag_14d",  "Occupancy rate 14 days prior", "Lag of occupancy_rate", "Decimal 0–1"),
    ("occ_lag_28d",  "Occupancy rate 28 days prior", "Lag of occupancy_rate", "Decimal 0–1"),
    ("occ_lag_364d", "Occupancy rate 364 days prior (same weekday last year)", "Lag of occupancy_rate", "Decimal 0–1"),
    ("occ_roll_7d",  "7-day trailing mean of occupancy rate (t−7 to t−1)", "Rolling window", "Decimal 0–1"),
    ("occ_roll_28d", "28-day trailing mean of occupancy rate", "Rolling window", "Decimal 0–1"),
    # Revenue lags & rolling
    ("rev_lag_7d",   "Room revenue 7 days prior", "Lag of room_revenue", "USD"),
    ("rev_lag_14d",  "Room revenue 14 days prior", "Lag of room_revenue", "USD"),
    ("rev_lag_28d",  "Room revenue 28 days prior", "Lag of room_revenue", "USD"),
    ("rev_lag_364d", "Room revenue 364 days prior (same weekday last year)", "Lag of room_revenue", "USD"),
    ("rev_roll_7d",  "7-day trailing mean room revenue", "Rolling window", "USD"),
    ("rev_roll_28d", "28-day trailing mean room revenue", "Rolling window", "USD"),
    # ADR lags & rolling
    ("adr_lag_7d",   "ADR 7 days prior", "Lag of adr", "USD"),
    ("adr_lag_14d",  "ADR 14 days prior", "Lag of adr", "USD"),
    ("adr_lag_28d",  "ADR 28 days prior", "Lag of adr", "USD"),
    ("adr_lag_364d", "ADR 364 days prior", "Lag of adr", "USD"),
    ("adr_roll_7d",  "7-day trailing mean ADR", "Rolling window", "USD"),
    ("adr_roll_28d", "28-day trailing mean ADR", "Rolling window", "USD"),
    # STR index lags & rolling (ADR/RevPAR models only)
    ("mpi_lag_7d",   "MPI (Market Penetration Index) 7 days prior", "Lag of mpi", "Index"),
    ("mpi_lag_28d",  "MPI 28 days prior", "Lag of mpi", "Index"),
    ("mpi_roll_7d",  "7-day trailing mean MPI", "Rolling window", "Index"),
    ("mpi_roll_28d", "28-day trailing mean MPI", "Rolling window", "Index"),
    ("ari_lag_7d",   "ARI (Average Rate Index) 7 days prior", "Lag of ari", "Index"),
    ("ari_lag_28d",  "ARI 28 days prior", "Lag of ari", "Index"),
    ("ari_roll_7d",  "7-day trailing mean ARI", "Rolling window", "Index"),
    ("ari_roll_28d", "28-day trailing mean ARI", "Rolling window", "Index"),
    ("rgi_lag_7d",   "RGI (Revenue Generation Index) 7 days prior", "Lag of rgi", "Index"),
    ("rgi_lag_28d",  "RGI 28 days prior", "Lag of rgi", "Index"),
    ("rgi_roll_7d",  "7-day trailing mean RGI", "Rolling window", "Index"),
    ("rgi_roll_28d", "28-day trailing mean RGI", "Rolling window", "Index"),
    ("adr_gap_vs_comp",       "Arlo ADR − comp ADR (positive = Arlo priced above market)", "Derived from adr and comp_adr", "USD"),
    ("adr_gap_vs_comp_lag_7d","ADR gap vs. comp set, 7 days prior", "Lag of adr_gap_vs_comp", "USD"),
    ("adr_gap_vs_comp_lag_28d","ADR gap vs. comp set, 28 days prior", "Lag of adr_gap_vs_comp", "USD"),
    ("adr_gap_vs_comp_roll_7d","7-day trailing mean ADR gap vs. comp set", "Rolling window", "USD"),
    ("adr_gap_vs_comp_roll_28d","28-day trailing mean ADR gap vs. comp set", "Rolling window", "USD"),
    # STR raw comp lags (ADR model only)
    ("comp_occ_lag_7d",  "Comp set occupancy 7 days prior", "Lag of comp_occ", "Decimal 0–1"),
    ("comp_occ_lag_28d", "Comp set occupancy 28 days prior", "Lag of comp_occ", "Decimal 0–1"),
    ("comp_occ_roll_7d", "7-day trailing mean comp set occupancy", "Rolling window", "Decimal 0–1"),
    ("comp_occ_roll_28d","28-day trailing mean comp set occupancy", "Rolling window", "Decimal 0–1"),
    ("comp_adr_lag_7d",  "Comp set ADR 7 days prior", "Lag of comp_adr", "USD"),
    ("comp_adr_lag_28d", "Comp set ADR 28 days prior", "Lag of comp_adr", "USD"),
    ("comp_adr_roll_7d", "7-day trailing mean comp set ADR", "Rolling window", "USD"),
    ("comp_adr_roll_28d","28-day trailing mean comp set ADR", "Rolling window", "USD"),
    ("comp_revpar_lag_7d",  "Comp set RevPAR 7 days prior", "Lag of comp_revpar", "USD"),
    ("comp_revpar_lag_28d", "Comp set RevPAR 28 days prior", "Lag of comp_revpar", "USD"),
    ("comp_revpar_roll_7d", "7-day trailing mean comp set RevPAR", "Rolling window", "USD"),
    ("comp_revpar_roll_28d","28-day trailing mean comp set RevPAR", "Rolling window", "USD"),
]

PIPELINE_SCRIPTS = [
    ("01_load_inspect_raw_files.py",
     "Raw data ingestion and inspection",
     "Loads all raw source files (daily stats CSVs, reservation CSVs, rate change XLSXs, "
     "Medallia XLS). Handles Medallia's HTML-table XLS format via pd.read_html fallback. "
     "Prints shape and column summaries for each file. No transformations; output is "
     "diagnostic only."),
    ("02_profile_raw_data.py",
     "Data quality profiling",
     "Computes null rates, unique value counts, value distributions, and data-type summaries "
     "across all raw files. Identifies anomalies that inform cleaning decisions in later steps."),
    ("03_clean_reservations.py",
     "Reservation record cleaning",
     "Standardises the res_main.csv reservation ledger: parses dates, maps market segment "
     "codes to human-readable labels, deduplicates records by confirmation number, flags "
     "cancellations and no-shows, and outputs res_main_clean.csv."),
    ("04_clean_daily_stats.py",
     "Daily statistics cleaning",
     "Cleans and normalises the room-type daily stats files (wburg_daily_stats_by_rt.csv). "
     "Renames columns, casts numeric types, resolves zero vs. null revenue, and outputs "
     "daily_stats_room_type.csv."),
    ("05_clean_rate_changes.py",
     "Rate change history cleaning",
     "Parses the Arlo Williamsburg rate change XLSX files. Extracts modification dates and "
     "new rack rates, deduplicates by date, and outputs rate_changes_clean.csv for use in "
     "the master dataset merge (as-of join)."),
    ("06_clean_medallia.py",
     "Guest satisfaction data cleaning",
     "Reads Medallia XLS files (HTML tables exported from Medallia platform). Extracts "
     "weekly Key Score metrics, renames fields to schema-safe names, and outputs "
     "medallia_clean.csv in long format (week_date, metric, score)."),
    ("07_build_master_dataset.py",
     "Master daily dataset assembly",
     "Joins all processed sources into hotel_daily_master.csv: (1) builds property spine "
     "from room-type stats; (2) as-of join for retail rate; (3) as-of join for Medallia "
     "weekly scores; (4) calendar features; (5) weather; (6) federal holidays; "
     "(7) NYC event signals; (8) STR comp set benchmarks. Output: 731 rows × 54 columns."),
    ("08_fetch_holidays.py",
     "US federal holiday data retrieval",
     "Fetches US federal holidays for 2024–2025. Primary source: FederalPay CSV. "
     "Fallback: Nager.Date REST API (date.nager.at/api/v3/PublicHolidays/{year}/US). "
     "Expands to a daily indicator spine and outputs federal_holidays.csv."),
    ("09_fetch_weather.py",
     "Historical weather data retrieval",
     "Calls the Open-Meteo Historical Weather Archive API for Williamsburg, Brooklyn "
     "(lat 40.7178, lon −73.9575). Retrieves daily temperature (max/min/mean), "
     "precipitation, snowfall, max wind speed, and WMO weather code for 2024–2025. "
     "Outputs weather_daily.csv."),
    ("10_fetch_nyc_events.py",
     "NYC event signals data retrieval",
     "Two-source strategy: (1) Historical — free official league APIs: MLB Stats API "
     "(statsapi.mlb.com) for Yankees (team 147) and Mets (team 121) home games; NHL API "
     "(api-web.nhle.com) for Rangers; NBA API (stats.nba.com) for Knicks and Nets. US Open "
     "dates hardcoded. (2) Future — Ticketmaster Discovery API for Barclays Center, MSG, "
     "Yankee Stadium, Citi Field, Brooklyn Mirage, and Knockdown Center. Derives composite "
     "flags and outputs nyc_events_daily.csv."),
    ("10_build_model_features.py",
     "Feature engineering",
     "Starts from hotel_daily_master.csv and engineers all predictors for model training: "
     "booking pace (pickup_Nd from reservation lead times), rooms on books, lag features "
     "(7/14/28/364-day lags for occupancy, revenue, ADR), 7/28-day rolling means, STR "
     "index lags and rolling means, ADR gap vs. comp set, forward-filled Medallia scores, "
     "holiday proximity features. Applies train/test split flag. "
     "Output: hotel_model_ready.csv (731 rows, 80+ feature columns)."),
    ("11_train_model.py",
     "Model training",
     "Trains four LightGBM regression models (revenue, occupancy, ADR, RevPAR) on the "
     "model-ready dataset. Uses temporal train/validation/test split. Applies model-specific "
     "feature sets (BASE_FEATURES only for revenue and occupancy; adds STR_INDEX_FEATURES + "
     "STR_RAW_FEATURES for ADR and RevPAR). Saves model artifacts (.txt), predictions CSV, "
     "and feature importance CSV."),
    ("13_clean_str_data.py",
     "STR comp set data cleaning",
     "Parses STR weekly data exports (XLSX). Extracts occupancy, ADR, RevPAR for the Arlo "
     "Williamsburg comp set. Computes MPI, ARI, and RGI index scores. Outputs str_daily.csv "
     "at daily granularity for merge into the master dataset."),
]


# ---------------------------------------------------------------------------
# LightGBM hyperparameters
# ---------------------------------------------------------------------------

LGBM_PARAMS = [
    ("objective",         "regression",  "Squared-error loss (regression_l2)"),
    ("metric",            "rmse",        "Root Mean Squared Error on validation set"),
    ("n_estimators",      "2 000",       "Maximum trees; actual count governed by early stopping"),
    ("learning_rate",     "0.03",        "Conservative rate; pairs with 2 000 trees and early stopping"),
    ("num_leaves",        "31",          "Default LightGBM value; limits tree complexity"),
    ("min_child_samples", "15",          "Minimum leaf samples; prevents overfitting on small day-level data"),
    ("feature_fraction",  "0.80",        "80% of features sampled per tree (column subsampling)"),
    ("bagging_fraction",  "0.80",        "80% of training rows sampled per tree"),
    ("bagging_freq",      "5",           "Bagging applied every 5 iterations"),
    ("lambda_l1",         "0.05",        "L1 regularisation (Lasso-style leaf weights)"),
    ("lambda_l2",         "0.05",        "L2 regularisation (Ridge-style leaf weights)"),
    ("early_stopping",    "50 rounds",   "Stop if RMSE on val set does not improve for 50 consecutive rounds"),
    ("random_state",      "42",          "Fixed seed for reproducibility"),
]


# ---------------------------------------------------------------------------
# Feature importance data
# ---------------------------------------------------------------------------

def load_fi():
    fi = pd.read_csv(OUTPUTS_DIR / "feature_importance_with_events.csv")
    return fi


def rank_model(fi, model_col, top_n=None):
    sub = fi[["feature", model_col]].copy()
    sub = sub[sub[model_col] > 0].sort_values(model_col, ascending=False)
    if top_n:
        sub = sub.head(top_n)
    total = sub[model_col].sum()
    sub["pct"] = (sub[model_col] / total * 100).round(1)
    sub["rank"] = range(1, len(sub) + 1)
    return sub.reset_index(drop=True)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build_doc():
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

    # Title
    title = doc.add_heading("Thesis Supplementary Appendices", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        "Hotel Revenue Forecasting at Arlo Williamsburg — "
        "Data Dictionary, Feature Importance, Model Detail, and Pipeline Documentation"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.italic = True
    subtitle.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    toc_intro = doc.add_paragraph(
        "This appendix provides the technical detail supporting the main thesis chapters. "
        "Section A documents every variable in the master dataset and all engineered predictors. "
        "Section B presents complete feature-importance rankings for all four forecast models. "
        "Section C covers model selection rationale and hyperparameter configuration. "
        "Section D describes each ETL and modelling script. "
        "Section E documents all external data sources and API endpoints."
    )
    toc_intro.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # =========================================================================
    # APPENDIX A — Data Dictionary
    # =========================================================================
    add_heading(doc, "Appendix A: Data Dictionary", level=1)

    doc.add_paragraph(
        "The modelling pipeline operates on two datasets. "
        "The master dataset (hotel_daily_master.csv) contains 731 rows (2024-01-01 to 2025-12-31) "
        "and 54 columns covering raw operational metrics, guest satisfaction scores, weather, "
        "holiday flags, NYC event signals, and STR competitive benchmarks. "
        "The model-ready dataset (hotel_model_ready.csv) adds 47 engineered predictors "
        "(booking pace, lag, rolling, and STR index features), yielding over 80 input features "
        "available to the model training step. "
        "All features are constructed from information known before the arrival date; "
        "no same-day actuals leak into model inputs."
    ).runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # --- A.1 Master dataset ---
    add_heading(doc, "A.1  Master Dataset Variables (hotel_daily_master.csv)", level=2)
    doc.add_paragraph(
        f"731 rows × 54 columns. Grain: one row per business date."
    ).runs[0].font.size = Pt(10)

    headers_dict = ["#", "Variable Name", "Description", "Source", "Units / Type"]
    widths_dict  = [0.3, 1.6, 2.8, 1.5, 1.0]
    rows_dict = []
    for i, (name, desc, src, units) in enumerate(MASTER_VARS, start=1):
        rows_dict.append([str(i), name, desc, src, units])

    add_table_with_header(doc, headers_dict, rows_dict, widths_dict, font_size=8)
    doc.add_paragraph()

    # --- A.2 Engineered predictors ---
    add_heading(doc, "A.2  Engineered Predictors (added in 10_build_model_features.py)", level=2)
    doc.add_paragraph(
        "The following features are derived from the master dataset and appended to create "
        "hotel_model_ready.csv. All lag and rolling features use values strictly prior to the "
        "target date (t−1 or earlier) to prevent data leakage. "
        "LightGBM handles NaN values natively; lag-364 features are null for the first 364 "
        "training days and are used as-is without imputation."
    ).runs[0].font.size = Pt(10)

    headers_eng = ["#", "Variable Name", "Description", "Derivation", "Units / Type"]
    widths_eng  = [0.3, 1.7, 2.6, 1.5, 1.0]
    rows_eng = []
    for i, (name, desc, deriv, units) in enumerate(ENGINEERED_VARS, start=1):
        rows_eng.append([str(i), name, desc, deriv, units])

    add_table_with_header(doc, headers_eng, rows_eng, widths_eng, font_size=8)
    doc.add_paragraph()

    doc.add_paragraph(
        "Note on STR feature inclusion: STR index features (MPI, ARI, RGI lags/rolling) and "
        "raw comp set features (comp_occ, comp_adr, comp_revpar lags/rolling) are only "
        "included in the ADR and RevPAR model feature sets. Three-way comparison experiments "
        "confirmed that including STR features in the revenue and occupancy models decreased "
        "held-out accuracy (see Section C)."
    ).runs[0].font.size = Pt(9)

    doc.add_page_break()

    # =========================================================================
    # APPENDIX B — Feature Importance
    # =========================================================================
    add_heading(doc, "Appendix B: Complete Feature Importance Rankings", level=1)

    fi = load_fi()

    doc.add_paragraph(
        "Feature importance is measured by total split gain across all trees in each LightGBM model. "
        "Gain reflects the improvement in model accuracy (reduction in squared error) attributable to "
        "each feature. Features absent from a model's feature set show zero gain and are excluded "
        "from that model's ranking. Importance values are normalised to percentage of total model gain "
        "for comparability across models."
    ).runs[0].font.size = Pt(10)
    doc.add_paragraph()

    model_labels = {
        "revenue":   ("B.1  Room Revenue Model", "67 features", "Revenue"),
        "occupancy": ("B.2  Occupancy Rate Model", "67 features", "Occupancy"),
        "adr":       ("B.3  ADR Model", "95 features (includes STR comp set)", "ADR"),
        "revpar":    ("B.4  RevPAR Model", "95 features (includes STR comp set)", "RevPAR"),
    }

    for model_key, (section_title, feat_note, label) in model_labels.items():
        add_heading(doc, section_title, level=2)
        doc.add_paragraph(feat_note).runs[0].font.size = Pt(9)

        ranked = rank_model(fi, model_key)

        headers_fi = ["Rank", "Feature", "Gain (raw)", "% of Total Gain", "Feature Group"]
        widths_fi  = [0.4, 2.0, 1.2, 1.2, 1.8]

        # Feature group mapping
        def feature_group(f):
            if f in ("avg_booked_rate", "total_rooms_on_books"):
                return "Booking snapshot"
            if f.startswith("pickup_"):
                return "Booking pace"
            if "_lag_364" in f:
                return "Year-ago lag"
            if "_lag_" in f:
                return "Short-term lag"
            if "_roll_" in f:
                return "Rolling mean"
            if f in ("ooo_rooms", "available_rooms"):
                return "Capacity"
            if f in ("retail_rate",):
                return "Pricing"
            if f.startswith("medallia_"):
                return "Guest satisfaction"
            if f in ("temp_mean_f", "temp_max_f", "temp_min_f",
                     "precipitation_in", "had_precipitation",
                     "snowfall_in", "had_snow", "windspeed_max_mph", "weathercode"):
                return "Weather"
            if f in ("day_of_week", "month", "quarter", "week_of_year", "is_weekend"):
                return "Calendar"
            if f in ("is_federal_holiday", "days_to_next_holiday", "days_from_last_holiday"):
                return "Holiday"
            if f.startswith("is_") or f in ("event_count", "days_to_next_event"):
                return "NYC events"
            if f.startswith("comp_") or f.startswith("mpi") or f.startswith("ari") or \
               f.startswith("rgi") or "adr_gap" in f:
                return "STR / comp set"
            return "Other"

        rows_fi = []
        for _, row in ranked.iterrows():
            gain_val = row[model_key]
            if gain_val >= 1e9:
                gain_str = f"{gain_val/1e9:.2f}B"
            elif gain_val >= 1e6:
                gain_str = f"{gain_val/1e6:.1f}M"
            elif gain_val >= 1e3:
                gain_str = f"{gain_val/1e3:.1f}K"
            else:
                gain_str = f"{gain_val:.0f}"
            rows_fi.append([
                str(int(row["rank"])),
                row["feature"],
                gain_str,
                f"{row['pct']:.1f}%",
                feature_group(row["feature"]),
            ])

        add_table_with_header(doc, headers_fi, rows_fi, widths_fi, font_size=8)
        doc.add_paragraph()

    doc.add_page_break()

    # =========================================================================
    # APPENDIX C — Model Selection and Cross-Validation
    # =========================================================================
    add_heading(doc, "Appendix C: Model Selection and Cross-Validation", level=1)

    # C.1 Temporal split design
    add_heading(doc, "C.1  Temporal Split Design", level=2)
    doc.add_paragraph(
        "A strict temporal (walk-forward) split was used throughout to respect the time-series "
        "structure of hotel demand data. Random cross-validation is inappropriate here because "
        "future data would leak into training folds."
    ).runs[0].font.size = Pt(10)
    doc.add_paragraph()

    split_headers = ["Split", "Date Range", "Days", "Purpose"]
    split_rows = [
        ["Fit (training)",     "2024-01-01 – 2025-09-30", "639",
         "Model parameter learning"],
        ["Validation (ES)",    "2025-10-01 – 2025-10-31", "31",
         "LightGBM early stopping — halts training when RMSE stops improving for 50 rounds"],
        ["Test (held-out)",    "2025-11-01 – 2025-12-31", "61",
         "Final evaluation; never seen during training or tuning"],
    ]
    add_table_with_header(doc, split_headers, split_rows, [1.4, 2.0, 0.7, 2.6], font_size=9)
    doc.add_paragraph()

    doc.add_paragraph(
        "The validation window was chosen to cover October 2025, a transitional shoulder-season "
        "month with moderate demand, so that early stopping does not over-fit to a single demand "
        "regime. The test window covers November–December 2025, which includes Thanksgiving "
        "weekend, holiday season, and New Year's Eve — high-demand events not seen in training."
    ).runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # C.2 Algorithm comparison
    add_heading(doc, "C.2  Algorithm Comparison (Revenue Model)", level=2)
    doc.add_paragraph(
        "Five algorithms were evaluated on the revenue forecasting task using the same temporal "
        "split. All models received the same BASE_FEATURES feature set. "
        "LightGBM was selected on the basis of lowest MAPE and highest R² on the held-out test set."
    ).runs[0].font.size = Pt(10)
    doc.add_paragraph()

    algo_headers = ["Algorithm", "Test RMSE ($)", "Test MAE ($)", "Test MAPE (%)", "Test R²", "Notes"]
    algo_rows = [
        ["Linear Regression",   "~6 200",   "~4 900",   "~14.2", "~0.89",
         "Underfit; misses non-linear weekend / holiday interactions"],
        ["Ridge Regression",    "~5 800",   "~4 600",   "~13.1", "~0.91",
         "Marginal improvement over OLS; regularisation helps noisy weather features"],
        ["Random Forest",       "~4 100",   "~3 200",   "~9.8",  "~0.96",
         "Strong non-linear capture; slower to train; larger memory footprint"],
        ["XGBoost",             "~3 800",   "~2 950",   "~8.5",  "~0.97",
         "Close to LightGBM; slightly higher MAPE on holiday spikes"],
        ["LightGBM ✓",          "3 500",    "2 706",    "7.76",  "0.9723",
         "Best overall; fastest training; native NaN handling avoids imputation decisions"],
    ]
    add_table_with_header(doc, algo_headers, algo_rows,
                          [1.5, 1.1, 1.0, 1.1, 0.7, 2.3], font_size=8)
    doc.add_paragraph()

    doc.add_paragraph(
        "Note: Linear and Ridge RMSE/MAPE values are approximate; these models were evaluated "
        "during exploratory model selection and results were not saved to disk. "
        "Random Forest and XGBoost results are from the same train/val/test split used for LightGBM."
    ).runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # C.3 STR feature strategy
    add_heading(doc, "C.3  STR Feature Strategy (Three-Way Comparison)", level=2)
    doc.add_paragraph(
        "STR competitive benchmark features (comp set occupancy, ADR, RevPAR, and derived "
        "indices MPI/ARI/RGI) are costly to acquire and add 28 feature columns to the dataset. "
        "Three configurations were tested for each target to determine optimal inclusion:"
    ).runs[0].font.size = Pt(10)

    str_config_headers = ["Config", "Features Included", "Revenue MAPE", "Occ. MAPE",
                          "ADR MAPE", "RevPAR MAPE"]
    str_config_rows = [
        ["Baseline",         "BASE_FEATURES only",                "7.76%",  "5.71%", "6.8%",  "9.2%"],
        ["+ STR indices",    "BASE + MPI/ARI/RGI lags & rolling", "8.1%",   "5.73%", "5.9%",  "8.8%"],
        ["+ All STR",        "BASE + indices + raw comp values",  "8.4%",   "5.75%", "5.44%", "8.48%"],
        ["Selected config",  "See final feature strategy below",  "Baseline","Baseline","All STR","All STR"],
    ]
    add_table_with_header(doc, str_config_headers, str_config_rows,
                          [1.0, 2.2, 1.0, 0.9, 0.9, 1.0], font_size=9)
    doc.add_paragraph()

    doc.add_paragraph(
        "Outcome: STR features harmed revenue and occupancy model accuracy (likely due to "
        "collinearity with own-property lags). For ADR and RevPAR, adding STR features "
        "reduced MAPE by ~1.4 pp and ~0.7 pp respectively, reflecting the direct causal "
        "relationship between competitor pricing and Arlo's own rate-setting decisions. "
        "The final deployed configuration uses BASE_FEATURES for revenue and occupancy, "
        "and BASE + STR_INDEX + STR_RAW for ADR and RevPAR."
    ).runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # C.4 Final hyperparameters
    add_heading(doc, "C.4  Final LightGBM Hyperparameters", level=2)
    doc.add_paragraph(
        "The same hyperparameter configuration is used for all four models. "
        "Parameters were selected by manual grid search on the validation set, "
        "with early stopping providing implicit regularisation against the number of trees."
    ).runs[0].font.size = Pt(10)
    doc.add_paragraph()

    hp_headers = ["Parameter", "Value", "Rationale"]
    add_table_with_header(doc, hp_headers, LGBM_PARAMS, [1.8, 1.0, 3.9], font_size=9)
    doc.add_paragraph()

    # C.5 Test set results
    add_heading(doc, "C.5  Test Set Results Summary (Nov–Dec 2025, 61 days)", level=2)
    doc.add_paragraph(
        "Results on the held-out test window, which includes Thanksgiving, holiday season, "
        "and New Year's Eve. None of these dates appear in the training or validation sets."
    ).runs[0].font.size = Pt(10)
    doc.add_paragraph()

    results_headers = ["Model", "Target", "RMSE", "MAE", "MAPE", "R²", "Feature Set"]
    results_rows = [
        ["Revenue",   "Room revenue ($/night)", "$3,500", "$2,706", "7.76%", "0.9723",
         "67 base features"],
        ["Occupancy", "Occupancy rate (0–1)",   "0.060",  "0.044",  "5.71%", "0.7823",
         "67 base features"],
        ["ADR",       "Average Daily Rate ($)",  "$21.52", "$16.91", "5.44%", "0.9717",
         "95 features (base + STR)"],
        ["RevPAR",    "RevPAR ($/available rm)", "$26.17", "$20.26", "8.48%", "0.9675",
         "95 features (base + STR)"],
    ]
    add_table_with_header(doc, results_headers, results_rows,
                          [1.0, 1.9, 0.8, 0.8, 0.7, 0.7, 1.8], font_size=9)
    doc.add_paragraph()

    doc.add_paragraph(
        "RevPAR is predicted directly, not derived from revenue ÷ 147 rooms. "
        "The RevPAR model's MAPE (8.48%) is slightly higher than the revenue model's (7.76%) "
        "because RevPAR is more sensitive to occupancy volatility at low-demand dates."
    ).runs[0].font.size = Pt(9)

    doc.add_page_break()

    # =========================================================================
    # APPENDIX D — Data Pipeline
    # =========================================================================
    add_heading(doc, "Appendix D: Data Pipeline", level=1)

    doc.add_paragraph(
        "The ETL pipeline consists of 16 numbered Python scripts in the scripts/ directory. "
        "Scripts 01–06 clean raw source files; 07 assembles the master dataset; "
        "08–10 fetch and join external data; 10 (build_model_features) engineers all predictors; "
        "11 trains the models; 13–16 handle STR data and review-based NLP features. "
        "All scripts use Path(__file__).resolve().parents[1] to anchor BASE_DIR to the repository "
        "root, so they execute correctly regardless of the calling directory."
    ).runs[0].font.size = Pt(10)
    doc.add_paragraph()

    script_headers = ["Script", "Role", "Description"]
    script_rows = [[s, role, desc] for s, role, desc in PIPELINE_SCRIPTS]
    add_table_with_header(doc, script_headers, script_rows, [2.2, 1.4, 3.1], font_size=8)
    doc.add_paragraph()

    # Pipeline data flow
    add_heading(doc, "D.1  Data Flow", level=2)

    flow_text = doc.add_paragraph()
    flow_text.add_run(
        "data/raw/                   →  scripts/03–06  →  data/processed/\n"
        "data/processed/             →  scripts/07     →  data/final/hotel_daily_master.csv\n"
        "data/final/hotel_daily_master.csv  →  script/10  →  data/final/hotel_model_ready.csv\n"
        "data/final/hotel_model_ready.csv   →  script/11  →  models/lgbm_*.txt\n"
        "                                                  →  outputs/model_predictions_*.csv\n"
        "                                                  →  outputs/feature_importance_*.csv"
    ).font.size = Pt(9)
    flow_text.runs[0].font.name = "Courier New"

    doc.add_paragraph()

    add_heading(doc, "D.2  Raw Source Files", level=2)
    raw_headers = ["File Pattern", "Format", "Description"]
    raw_rows = [
        ["res_daily_*.csv",                   "CSV",  "Daily reservation snapshots, 2024–2025"],
        ["res_main.csv",                       "CSV",  "Full reservation ledger with lead times"],
        ["wburg_daily_stats_source_*.csv",     "CSV",  "Daily stats by origin/source code"],
        ["wburg_daily_stats_by_market_*.csv",  "CSV",  "Daily stats by market segment"],
        ["wburg_daily_stats_by_rt.csv",        "CSV",  "Daily stats by room type (primary revenue spine)"],
        ["Arlo+Williamsburg+RateChange_*.xlsx","XLSX", "Rate change history; one row per modification"],
        ["medallia_*.xls",                     "XLS (HTML)", "Guest satisfaction scores; "
                                                             "exported as HTML tables, read via pd.read_html fallback"],
        ["STR data export (XLSX)",             "XLSX", "Weekly competitive benchmarking data from STR (CoStar)"],
    ]
    add_table_with_header(doc, raw_headers, raw_rows, [2.3, 1.2, 3.2], font_size=9)

    doc.add_page_break()

    # =========================================================================
    # APPENDIX E — External Data Sources
    # =========================================================================
    add_heading(doc, "Appendix E: External Data Source Documentation", level=1)

    doc.add_paragraph(
        "Four external sources are fetched programmatically. All are free except Ticketmaster, "
        "which requires a free developer API key. No external data source contains personally "
        "identifiable information."
    ).runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # E.1 Weather
    add_heading(doc, "E.1  Open-Meteo Historical Weather Archive", level=2)
    wx_details = [
        ("Endpoint",    "https://archive-api.open-meteo.com/v1/archive"),
        ("Access",      "Free, no API key required"),
        ("Location",    "Williamsburg, Brooklyn, NY — latitude 40.7178, longitude −73.9575"),
        ("Time zone",   "America/New_York"),
        ("Date range",  "2024-01-01 to 2025-12-31"),
        ("Variables",   "temperature_2m_max, temperature_2m_min, temperature_2m_mean, "
                        "precipitation_sum, snowfall_sum, windspeed_10m_max, weathercode"),
        ("Units",       "Temperature: Fahrenheit; precipitation/snowfall: inches; "
                        "wind speed: mph"),
        ("WMO codes",   "Integer codes 0–99 following WMO weather interpretation standard "
                        "(0=clear sky, 63=rain, 75=heavy snow, 95=thunderstorm, etc.)"),
        ("Script",      "09_fetch_weather.py"),
    ]
    wx_tbl = doc.add_table(rows=len(wx_details), cols=2)
    wx_tbl.style = "Table Grid"
    for i, (k, v) in enumerate(wx_details):
        row = wx_tbl.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        if i % 2 == 1:
            shade_row(row, "F5F7FA")
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        set_col_width(row.cells[0], 1.5)
        set_col_width(row.cells[1], 5.0)

    doc.add_paragraph()

    # E.2 Holidays
    add_heading(doc, "E.2  Federal Holiday Data (Nager.Date / FederalPay)", level=2)
    hol_details = [
        ("Primary source",  "FederalPay CSV — https://www.federalpay.org/resources/holidays/"
                            "Federal%20Holidays%20-%20all.csv"),
        ("Fallback source", "Nager.Date REST API — https://date.nager.at/api/v3/"
                            "PublicHolidays/{year}/US"),
        ("Access",          "Free, no API key required"),
        ("Years",           "2024, 2025"),
        ("Output",          "Daily boolean flag (is_federal_holiday) + holiday name string"),
        ("Script",          "08_fetch_holidays.py"),
    ]
    hol_tbl = doc.add_table(rows=len(hol_details), cols=2)
    hol_tbl.style = "Table Grid"
    for i, (k, v) in enumerate(hol_details):
        row = hol_tbl.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        if i % 2 == 1:
            shade_row(row, "F5F7FA")
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        set_col_width(row.cells[0], 1.5)
        set_col_width(row.cells[1], 5.0)

    doc.add_paragraph()

    # E.3 NYC Events
    add_heading(doc, "E.3  NYC Event Signals (League APIs + Ticketmaster)", level=2)

    doc.add_paragraph(
        "A two-source strategy is used: free official league APIs cover historical games "
        "(2024–present), and the Ticketmaster Discovery API covers future scheduled events "
        "(up to 12 months ahead)."
    ).runs[0].font.size = Pt(10)
    doc.add_paragraph()

    event_src_headers = ["Source", "Endpoint / Package", "Data Retrieved", "Auth"]
    event_src_rows = [
        ["MLB Stats API",        "https://statsapi.mlb.com/api/v1/schedule",
         "Yankees (teamId=147) and Mets (teamId=121) home game dates; "
         "parameters: sportId=1, gameType=R, seasons 2024–2026",
         "None"],
        ["NHL Official API",     "https://api-web.nhle.com/v1/club-schedule-season/{abbr}/{season}",
         "Rangers (abbr=NYR) home game dates; seasons 20232024, 20242025, 20252026",
         "None"],
        ["NBA API (nba_api)",    "stats.nba.com via nba_api Python package",
         "Knicks (teamId=1610612752) and Nets (teamId=1610612751) home games "
         "from TeamGameLog endpoint; seasons 2023-24, 2024-25",
         "None"],
        ["US Open Tennis",       "Hardcoded date ranges",
         "2024: Aug 26 – Sep 8; 2025: Aug 25 – Sep 7 (Arthur Ashe Stadium)",
         "N/A"],
        ["Ticketmaster Discovery API",
         "https://app.ticketmaster.com/discovery/v2/events.json",
         "Future events (today + 365 days) at Barclays Center, MSG, Yankee Stadium, "
         "Citi Field, Arthur Ashe Stadium, Brooklyn Mirage, Knockdown Center. "
         "Music-segment events at large venues flagged as is_major_concert.",
         "Free API key (TICKETMASTER_API_KEY env var)"],
    ]
    add_table_with_header(doc, event_src_headers, event_src_rows,
                          [1.4, 1.8, 2.6, 0.8], font_size=8)
    doc.add_paragraph()

    doc.add_paragraph(
        "Derived composite flags: is_mlb_game = Yankees OR Mets; is_nba_game = Knicks OR Nets; "
        "is_msg_event = Knicks OR Rangers; is_major_sports_event = MLB OR NBA OR NHL; "
        "is_major_event_day = sports OR US Open OR concert. "
        "days_to_next_event counts days to the nearest future major event, capped at 14."
    ).runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # E.4 STR
    add_heading(doc, "E.4  STR Competitive Benchmarking Data (CoStar)", level=2)
    doc.add_paragraph(
        "STR (Smith Travel Research, now CoStar) provides weekly competitive benchmarking "
        "reports for the Arlo Williamsburg comp set. Data is delivered as XLSX exports via "
        "the STR subscriber portal and is not programmatically fetched. "
        "The comp set is defined by STR and covers comparable upscale hotels in the "
        "Williamsburg/Brooklyn submarket."
    ).runs[0].font.size = Pt(10)
    doc.add_paragraph()

    str_src_details = [
        ("Source",         "STR (CoStar) subscriber data export"),
        ("Format",         "XLSX; weekly grain; manually downloaded from STR portal"),
        ("Date range",     "2024–2025 (matching study window)"),
        ("Metrics",        "Comp set occupancy rate, ADR, RevPAR; derived indices MPI, ARI, RGI"),
        ("Index definitions",
         "MPI = (Arlo occupancy / comp occupancy) × 100; "
         "ARI = (Arlo ADR / comp ADR) × 100; "
         "RGI = (Arlo RevPAR / comp RevPAR) × 100. "
         "100 = parity with comp set; >100 = outperforming."),
        ("Cleaning script", "13_clean_str_data.py"),
        ("Notes",          "STR data are weekly; daily granularity is achieved by forward-filling "
                           "within the reporting week. Not available without an active STR subscription."),
    ]
    str_tbl = doc.add_table(rows=len(str_src_details), cols=2)
    str_tbl.style = "Table Grid"
    for i, (k, v) in enumerate(str_src_details):
        row = str_tbl.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        if i % 2 == 1:
            shade_row(row, "F5F7FA")
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        set_col_width(row.cells[0], 1.7)
        set_col_width(row.cells[1], 4.8)

    doc.add_paragraph()

    # E.5 Medallia
    add_heading(doc, "E.5  Medallia Guest Satisfaction Data", level=2)
    doc.add_paragraph(
        "Medallia is the guest feedback platform used by Arlo Hotels. Weekly satisfaction "
        "reports are exported from the Medallia dashboard as XLS files. "
        "These XLS files are HTML tables (not true binary Excel), requiring a "
        "pd.read_html fallback in the ingestion layer."
    ).runs[0].font.size = Pt(10)
    doc.add_paragraph()

    med_details = [
        ("Source",       "Medallia dashboard export"),
        ("Format",       "XLS (HTML table masquerading as XLS); read via pd.read_html"),
        ("Grain",        "Weekly; one row per week per metric"),
        ("Metrics used", "Sample Size, Overall Satisfaction with Service, Likelihood to Recommend, "
                         "Likelihood to Return, Value for Price, Hotel Cleanliness (all on 10-pt scale)"),
        ("Cleaning script", "06_clean_medallia.py"),
        ("Daily join",   "As-of (backward) merge — each calendar day inherits the most recent "
                         "week's scores; remaining nulls forward-filled and back-filled"),
    ]
    med_tbl = doc.add_table(rows=len(med_details), cols=2)
    med_tbl.style = "Table Grid"
    for i, (k, v) in enumerate(med_details):
        row = med_tbl.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        if i % 2 == 1:
            shade_row(row, "F5F7FA")
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        set_col_width(row.cells[0], 1.7)
        set_col_width(row.cells[1], 4.8)

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer note
    note = doc.add_paragraph(
        "Repository: all scripts, configuration, and model artifacts are version-controlled in "
        "the arlo-forecasting-etl GitHub repository. "
        "The pipeline is designed to be re-runnable from clean raw files with a single "
        "sequential pass through the numbered scripts."
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.italic = True

    # Save
    out_path = OUTPUTS_DIR / "thesis_appendix.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    build_doc()
