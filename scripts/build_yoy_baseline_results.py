"""
Build yoy_baseline_results.docx — YoY Baseline Results section for thesis.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUTS_DIR = Path(__file__).resolve().parents[1] / "outputs"

NAVY  = "1F3864"
LIGHT = "DCE6F1"
ALT   = "F2F7FB"
WHITE = "FFFFFF"
RED   = "C00000"
GREEN = "375623"


def _shd(cell, hex_fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _col_width(cell, inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    w = OxmlElement("w:tcW")
    w.set(qn("w:w"), str(int(inches * 1440)))
    w.set(qn("w:type"), "dxa")
    tcPr.append(w)


def _cell(cell, text, bold=False, size=10.5, color=None,
          align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def heading(doc, text, level=1, color="1F3864", size=None):
    h = doc.add_heading(text, level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sizes = {1: 15, 2: 13, 3: 11}
    for run in h.runs:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
        if size:
            run.font.size = Pt(size)
        elif level in sizes:
            run.font.size = Pt(sizes[level])
    return h


def body(doc, text, size=11, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead)
        r.font.bold = True
        r.font.size = Pt(size)
    r2 = p.add_run(text)
    r2.font.size = Pt(size)
    return p


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.25)
    sec.right_margin  = Inches(1.25)

    # ── Title ──────────────────────────────────────────────────────────────
    heading(doc, "4.2  Year-over-Year Baseline Results", level=1)
    doc.add_paragraph()

    body(doc,
         "Before evaluating the machine learning models, an experience-based year-over-year "
         "(YoY) forecast was constructed and evaluated on the same 61-day holdout period "
         "(November 1 – December 31, 2025). This baseline represents the most defensible "
         "manual forecasting method available to a hotel revenue manager: one who has access "
         "to the prior year's actual performance and applies an informed trend adjustment based "
         "on how the current year has been tracking. Its purpose is to establish a performance "
         "ceiling for human-driven forecasting methods against which the ML model can be "
         "objectively compared.")

    doc.add_paragraph()

    # ── 4.2.1 Baseline Construction ─────────────────────────────────────────
    heading(doc, "4.2.1  How the Baseline Was Constructed", level=2)
    doc.add_paragraph()

    body(doc,
         "The experience-based YoY forecast was built in three stages. First, the actual "
         "daily values for each of the four target KPIs — room revenue, occupancy rate, "
         "average daily rate (ADR), and RevPAR — were extracted from the PMS daily statistics "
         "for November 1 – December 31, 2024. These 61 days served as the raw prior-year "
         "base. Second, the year-to-date performance from January through October was compared "
         "across 2024 and 2025 to compute a KPI-specific growth factor for each metric:")

    doc.add_paragraph()

    # Growth factor table
    gf_headers = ["KPI", "Jan–Oct 2024 Daily Mean", "Jan–Oct 2025 Daily Mean", "YTD Growth Factor"]
    gf_widths  = [1.2, 1.8, 1.8, 1.7]
    gf_data = [
        ("Room Revenue",   "$38,165 / day",  "$41,304 / day",  "+8.2%"),
        ("Occupancy Rate", "86.3%",           "87.8%",           "+1.8%"),
        ("ADR",            "$314.90",         "$320.20",         "+1.7%"),
        ("RevPAR",         "$278.64",         "$288.79",         "+3.7%"),
    ]

    gf_table = doc.add_table(rows=1, cols=4)
    gf_table.style = "Table Grid"
    gf_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = gf_table.rows[0]
    for i, (h, w) in enumerate(zip(gf_headers, gf_widths)):
        _shd(hrow.cells[i], NAVY)
        _cell(hrow.cells[i], h, bold=True, size=10, color=WHITE,
              align=WD_ALIGN_PARAGRAPH.CENTER)
        _col_width(hrow.cells[i], w)

    for ri, row_data in enumerate(gf_data):
        row = gf_table.add_row()
        fill = ALT if ri % 2 else WHITE
        for ci, (cell, val) in enumerate(zip(row.cells, row_data)):
            _shd(cell, fill)
            _cell(cell, val, size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT)
            _col_width(cell, gf_widths[ci])

    doc.add_paragraph()

    cap1 = doc.add_paragraph(
        "Table B1: YTD growth factors derived from Jan–Oct 2025 vs. Jan–Oct 2024 actuals, "
        "used to adjust prior-year base values for the experienced YoY forecast."
    )
    cap1.runs[0].font.size = Pt(9)
    cap1.runs[0].font.italic = True

    doc.add_paragraph()

    body(doc,
         "Third, each daily value from November–December 2024 was multiplied by its "
         "corresponding KPI growth factor to produce the 2025 forecast. A 364-day offset "
         "(exactly 52 weeks) was used rather than a simple calendar-date match, ensuring "
         "day-of-week alignment between the base year and the forecast period — the Saturday "
         "before Thanksgiving 2025 maps to the Saturday before Thanksgiving 2024, not to "
         "an arbitrary weekday. The resulting formula for each forecast date is:")

    doc.add_paragraph()

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = formula.add_run(
        "Forecast(t, KPI)  =  Actual(t − 364 days, KPI)  ×  (1 + YTD Growth Factorₖᵖᴵ)"
    )
    r_f.font.size = Pt(11)
    r_f.font.bold = True
    r_f.font.name = "Courier New"

    doc.add_paragraph()

    body(doc,
         "This approach mirrors the reasoning an experienced revenue manager would apply: "
         "anchor the forecast in prior-year actuals, which encode known seasonal patterns, "
         "weekday effects, and hotel-specific demand cycles, then scale by the observed "
         "growth trajectory of the business. It is considerably more sophisticated than "
         "a naive same-period replication because it accounts for structural shifts in "
         "demand level across years.")

    doc.add_paragraph()

    # ── 4.2.2 Results by KPI ────────────────────────────────────────────────
    heading(doc, "4.2.2  Baseline Performance by KPI", level=2)
    doc.add_paragraph()

    body(doc,
         "The experience-based YoY forecast was evaluated against the 61-day held-out "
         "actuals using four metrics: MAPE, MAE, RMSE, and R². Results are presented "
         "in Table B2 and discussed individually below.")

    doc.add_paragraph()

    # Results table
    res_headers = ["KPI", "MAPE", "MAE", "RMSE", "R²", "Interpretation"]
    res_widths  = [1.1, 0.7, 1.0, 1.1, 0.6, 2.6]
    res_data = [
        ("Room Revenue",   "21.10%", "$8,788",  "$12,729", "0.63",
         "Moderate predictive power; prior-year revenue patterns partially transfer but "
         "day-level spikes are missed"),
        ("Occupancy Rate", "15.54%", "13.2 pp",  "17.7 pp", "−0.96",
         "R² below zero — the YoY projection is less accurate than predicting the "
         "period mean; short-term booking behaviour is not captured"),
        ("ADR",            "15.10%", "$49.00",  "$65.04",  "0.74",
         "Best-performing KPI for YoY; rate levels partially anchor to prior year, "
         "but competitive pricing dynamics add unexplained variance"),
        ("RevPAR",         "21.22%", "$62.11",  "$88.93",  "0.62",
         "Reflects compounded error from both occupancy and ADR projections; "
         "limited usefulness as a pricing anchor"),
    ]

    res_table = doc.add_table(rows=1, cols=6)
    res_table.style = "Table Grid"
    res_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow2 = res_table.rows[0]
    for i, (h, w) in enumerate(zip(res_headers, res_widths)):
        _shd(hrow2.cells[i], NAVY)
        _cell(hrow2.cells[i], h, bold=True, size=10, color=WHITE,
              align=WD_ALIGN_PARAGRAPH.CENTER)
        _col_width(hrow2.cells[i], w)

    for ri, row_data in enumerate(res_data):
        row = res_table.add_row()
        fill = ALT if ri % 2 else WHITE
        for ci, (cell, val) in enumerate(zip(row.cells, row_data)):
            _shd(cell, fill)
            align = WD_ALIGN_PARAGRAPH.LEFT if ci in (0, 5) else WD_ALIGN_PARAGRAPH.CENTER
            _cell(cell, val, size=10, align=align)
            _col_width(cell, res_widths[ci])

    doc.add_paragraph()
    cap2 = doc.add_paragraph(
        "Table B2: Experience-based YoY forecast performance on the 61-day holdout "
        "(Nov–Dec 2025). pp = percentage points."
    )
    cap2.runs[0].font.size = Pt(9)
    cap2.runs[0].font.italic = True

    doc.add_paragraph()
    doc.add_page_break()

    # ── KPI narratives ───────────────────────────────────────────────────────
    heading(doc, "4.2.3  KPI-Level Analysis", level=2)
    doc.add_paragraph()

    kpi_blocks = [
        (
            "Room Revenue",
            "21.10% MAPE  |  MAE $8,788  |  R² = 0.63",
            "The YoY revenue forecast achieved a MAPE of 21.10% and an R² of 0.63 across the "
            "61-day test window. While the R² indicates that prior-year revenue patterns explain "
            "roughly 63% of the day-level variance in 2025 revenue, the remaining 37% — "
            "corresponding to average daily errors of $8,788 — represents a material shortfall "
            "for operational planning. Errors were concentrated on days where 2025 demand "
            "diverged structurally from 2024: Thanksgiving week (shifted by one calendar day "
            "relative to 2024), the first week of December (which saw an unusually strong "
            "corporate group in 2025 with no 2024 equivalent), and the final week of December "
            "where New Year's Eve demand was materially higher in 2025. The flat growth "
            "multiplier of +8.2% — derived from the January–October YTD trend — could not "
            "distinguish between high-demand and low-demand days within the test window; it "
            "scaled all days proportionally regardless of their individual demand drivers."
        ),
        (
            "Occupancy Rate",
            "15.54% MAPE  |  MAE 13.2 pp  |  R² = −0.96",
            "The occupancy forecast produced the weakest result of the four KPIs, with a "
            "MAPE of 15.54% and a negative R² of −0.96. A negative R² indicates that the "
            "YoY model performed worse than simply predicting the average occupancy rate for "
            "the period — a strong signal that prior-year occupancy patterns are an unreliable "
            "guide to day-level demand in the holdout window. This result is consistent with "
            "the nature of hotel occupancy: it is driven primarily by short-term booking "
            "behaviour (arrivals in the final 7–30 days before a stay), cancellation and "
            "no-show patterns, and group block activity — all of which change from year to "
            "year and cannot be inferred from a 12-month-old observation. The +1.8% growth "
            "multiplier applied to 2024 occupancy levels amplified rather than corrected "
            "these mismatches on days where 2025 occupancy moved in the opposite direction "
            "to 2024. An experienced revenue manager reviewing these results would recognise "
            "that occupancy forecasting requires real-time booking pace data, which the "
            "YoY method structurally cannot provide."
        ),
        (
            "ADR",
            "15.10% MAPE  |  MAE $49.00  |  R² = 0.74",
            "ADR is the KPI where the experience-based approach performs best, achieving an "
            "R² of 0.74 and a MAPE of 15.10%. This reflects the relative stability of hotel "
            "rack rates year-over-year: published rate structures, seasonal pricing tiers, and "
            "negotiated group rates tend to persist across years with moderate inflation "
            "adjustments. The +1.7% YTD growth factor applied to 2024 ADR levels captured "
            "this general upward drift reasonably well for mid-range demand days. However, "
            "the MAPE of 15.10% and MAE of $49.00 per night reveal that the YoY method "
            "cannot account for competitive rate dynamics — periods when Arlo's pricing "
            "diverged from the prior year due to competitor rate changes, demand compression "
            "events, or deliberate yield management decisions. These factors are invisible to "
            "a backward-looking YoY trend and represent the primary driver of ADR forecast "
            "error in the test window."
        ),
        (
            "RevPAR",
            "21.22% MAPE  |  MAE $62.11  |  R² = 0.62",
            "RevPAR combines the pricing signal of ADR with the volume signal of occupancy, "
            "making it the most comprehensive indicator of revenue productivity per available "
            "room. The YoY forecast achieved a MAPE of 21.22% and R² of 0.62 for RevPAR — "
            "a result that closely tracks the revenue model's performance and is notably "
            "worse than the ADR model alone. This gap illustrates the compounding effect of "
            "occupancy forecast error: even when the rate component is reasonably estimated, "
            "the volume error propagates through to RevPAR. The MAPE of 21.22% translates "
            "to an average daily error of $62.11 per available room, or approximately $9,130 "
            "per day in absolute revenue terms for a 147-room property — a level of "
            "imprecision that would materially affect pricing, staffing, and inventory "
            "decisions if used as an operational input. The R² of 0.62, while positive, "
            "confirms that the YoY approach leaves substantial day-level variance "
            "unexplained, setting a clear performance floor for the machine learning "
            "evaluation that follows."
        ),
    ]

    for kpi_name, subtitle, text in kpi_blocks:
        heading(doc, kpi_name, level=3, color="2E74B5")
        sub = doc.add_paragraph(subtitle)
        sub.runs[0].font.size = Pt(10)
        sub.runs[0].font.italic = True
        sub.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        doc.add_paragraph()
        body(doc, text)
        doc.add_paragraph()

    # ── 4.2.4 Summary ──────────────────────────────────────────────────────
    heading(doc, "4.2.4  Summary of Baseline Limitations", level=2)
    doc.add_paragraph()

    body(doc,
         "Taken together, the YoY baseline results reveal three structural limitations of "
         "experience-based forecasting that motivate the machine learning approach:")

    doc.add_paragraph()

    limitations = [
        ("No real-time demand signal.",
         "The YoY method has no access to current booking pace, rooms on books, or "
         "pickup trends. These signals — which update daily and are strong predictors of "
         "near-term demand — are the primary source of forecast accuracy that the ML "
         "model exploits. The occupancy R² of −0.96 is the clearest evidence of this gap."),

        ("Flat growth multiplier cannot capture day-level heterogeneity.",
         "Applying a single YTD growth factor to every day in the forecast window assumes "
         "that 2025 demand grew uniformly over 2024 across all days and demand patterns. "
         "In practice, high-demand weekends, holiday windows, and event days grew at very "
         "different rates than mid-week shoulder days. The YoY method cannot distinguish "
         "these without a day-level adjustment model — which would itself require the kind "
         "of feature engineering used in the ML pipeline."),

        ("No competitive or external signal.",
         "The method is entirely backward-looking. It cannot incorporate information about "
         "competitor pricing (STR comp set data), upcoming NYC events, weather forecasts, "
         "or any other forward-looking demand driver. As the ADR results show, this "
         "limitation is most consequential for rate-setting decisions, where competitive "
         "context materially affects optimal pricing."),
    ]

    for i, (lead, text) in enumerate(limitations, 1):
        p = doc.add_paragraph(style="List Number")
        r1 = p.add_run(f"{lead}  ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)

    doc.add_paragraph()

    body(doc,
         "These limitations do not invalidate the experience-based YoY approach as a "
         "practical tool — it remains widely used in hotel revenue management precisely "
         "because it is fast, transparent, and requires no data infrastructure. Rather, "
         "they define the performance ceiling of human-driven forecasting on this dataset "
         "and establish the benchmark against which the LightGBM models are evaluated in "
         "Section 4.3.")

    # ── Save ────────────────────────────────────────────────────────────────
    out = OUTPUTS_DIR / "yoy_baseline_results.docx"
    doc.save(str(out))
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
